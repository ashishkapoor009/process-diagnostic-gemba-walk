"""Word (.docx) report generation via python-docx, mirroring the PDF
report's structure for stakeholders who prefer an editable document.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.reports.report_data import ReportContext

BRAND_BLUE = RGBColor(0x1D, 0x4E, 0xD8)


def _heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
    return h


def _add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def generate_word_report(ctx: ReportContext) -> bytes:
    doc = Document()

    title = doc.add_heading("Process Diagnostic / Gemba Walk Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(ctx.metadata.process_name).runs[0].font.size = Pt(16)
    doc.add_paragraph(
        f"{ctx.metadata.team_name} | {ctx.metadata.lob} | "
        f"Generated {ctx.generated_at:%Y-%m-%d %H:%M} UTC"
    )

    _heading(doc, "Executive Summary")
    for para in ctx.executive_summary.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_page_break()
    _heading(doc, "Process Overview")
    _add_table(
        doc, ["Metric", "Value"],
        [
            ["Current FTE", ctx.metadata.current_fte],
            ["Current Volume / Period", ctx.metadata.current_volume],
            ["Average Handle Time (min)", ctx.metadata.aht_minutes],
            ["Pain Areas", ctx.metadata.pain_areas or "Not stated"],
            ["Dependencies", ctx.metadata.dependencies or "Not stated"],
            ["Known Risks", ctx.metadata.known_risks or "Not stated"],
            ["Compliance Requirements", ctx.metadata.compliance_requirements or "Not stated"],
        ],
    )

    doc.add_page_break()
    _heading(doc, "Current State - Process Diagnostics")
    _add_table(
        doc, ["#", "Step", "Owner", "VA/NVA/BNVA", "Cycle(m)", "Wastes", "Root Cause", "Automation", "AI Ready"],
        [
            [
                d.step_number, d.step_name, d.owner, d.value_classification.value, d.cycle_time_minutes,
                ", ".join(w.value for w in d.lean_wastes) or "None", d.root_cause or "-",
                d.automation_score, d.ai_readiness_score,
            ]
            for d in ctx.diagnostics
        ],
    )

    for title_text, recs in [
        ("Lean & Standardization Findings", ctx.lean_recommendations),
        ("Automation Findings", ctx.automation_recommendations),
        ("AI / GenAI Findings", ctx.ai_recommendations),
    ]:
        doc.add_page_break()
        _heading(doc, title_text)
        if not recs:
            doc.add_paragraph("No recommendations in this category.")
        for r in recs:
            doc.add_heading(f"{r.title} ({r.category.value})", level=3)
            doc.add_paragraph(r.description)
            doc.add_paragraph(f"Rationale: {r.rationale}")
            doc.add_paragraph(
                f"Impact: {r.prioritization.business_impact}/10 | Effort: {r.prioritization.implementation_effort}/10 | "
                f"ROI: {r.prioritization.roi}/10 | Horizon: {r.roadmap_horizon.value} | "
                f"FTE Savings: {r.savings.fte_savings} | Annual Savings: ${r.savings.annual_cost_savings:,.0f} | "
                f"Confidence: {r.confidence_score:.0%} ({r.source_type.value})"
            )

    doc.add_page_break()
    _heading(doc, "Quick Wins")
    _add_table(
        doc, ["Title", "Category", "Impact", "Effort", "Timeline"],
        [[r.title, r.category.value, r.prioritization.business_impact, r.prioritization.implementation_effort,
           r.roadmap_horizon.value] for r in ctx.quick_wins],
    )

    doc.add_page_break()
    _heading(doc, "Implementation Roadmap")
    for horizon, recs in ctx.roadmap_grouped().items():
        doc.add_heading(horizon, level=3)
        for r in recs:
            doc.add_paragraph(f"{r.title} ({r.category.value})", style="List Bullet")

    doc.add_page_break()
    _heading(doc, "Savings & Efficiency Summary")
    s = ctx.savings_summary
    _add_table(
        doc, ["Metric", "Value"],
        [
            ["Total Recommendations", s.get("total_recommendations", "-")],
            ["Estimated FTE Savings (FTEs Released)", s.get("total_fte_savings", "-")],
            ["In-Year Savings", f"${s.get('in_year_savings', 0):,.0f}"],
            ["12-Month Savings (Full Run-Rate)", f"${s.get('twelve_month_savings', 0):,.0f}"],
            ["Blended Efficiency Improvement", f"{s.get('blended_efficiency_improvement_pct', '-')}%"],
            ["Target Efficiency Range", s.get("target_efficiency_range_pct", "25-30%")],
            ["Annual FTE Cost (User Provided)", f"${s.get('annual_fte_cost', 0):,.0f}"],
        ],
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
